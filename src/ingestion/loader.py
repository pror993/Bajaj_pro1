import os
import yaml
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import email
from email import policy
import re


class DocumentLoader:
    def __init__(self, config_path: str = "ingestion_config.yaml"):
        """Initialize document loader with configuration."""
        self.config = self._load_config(config_path)
        self.supported_formats = self.config.get("ingestion", {}).get("supported_formats", [])
    
    def _load_config(self, config_path: str) -> Dict:
        """Load ingestion configuration."""
        try:
            with open(config_path, 'r') as file:
                return yaml.safe_load(file)
        except Exception as e:
            print(f"Error loading config: {e}")
            return {}
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load document based on file extension.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing document content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._load_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._load_word(file_path)
        elif file_extension == '.txt':
            return self._load_text(file_path)
        elif file_extension == '.eml':
            return self._load_email(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Load PDF document."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    content += f"\n--- Page {page_num + 1} ---\n"
                    content += page.extract_text()
                
                return {
                    'content': content,
                    'metadata': {
                        'file_path': str(file_path),
                        'file_type': 'pdf',
                        'num_pages': len(pdf_reader.pages),
                        'file_size': file_path.stat().st_size
                    }
                }
        except Exception as e:
            raise Exception(f"Error loading PDF {file_path}: {e}")
    
    def _load_word(self, file_path: Path) -> Dict[str, Any]:
        """Load Word document."""
        try:
            doc = Document(file_path)
            
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
            
            return {
                'content': content,
                'metadata': {
                    'file_path': str(file_path),
                    'file_type': 'word',
                    'num_paragraphs': len(doc.paragraphs),
                    'num_tables': len(doc.tables),
                    'file_size': file_path.stat().st_size
                }
            }
        except Exception as e:
            raise Exception(f"Error loading Word document {file_path}: {e}")
    
    def _load_text(self, file_path: Path) -> Dict[str, Any]:
        """Load text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                'content': content,
                'metadata': {
                    'file_path': str(file_path),
                    'file_type': 'text',
                    'file_size': file_path.stat().st_size
                }
            }
        except Exception as e:
            raise Exception(f"Error loading text file {file_path}: {e}")
    
    def _load_email(self, file_path: Path) -> Dict[str, Any]:
        """Load email file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                email_content = file.read()
            
            # Parse email
            msg = email.message_from_string(email_content, policy=policy.default)
            
            # Extract email parts
            subject = msg.get('subject', '')
            sender = msg.get('from', '')
            date = msg.get('date', '')
            
            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body += part.get_content()
            else:
                body = msg.get_content()
            
            return {
                'content': f"Subject: {subject}\nFrom: {sender}\nDate: {date}\n\n{body}",
                'metadata': {
                    'file_path': str(file_path),
                    'file_type': 'email',
                    'subject': subject,
                    'sender': sender,
                    'date': date,
                    'file_size': file_path.stat().st_size
                }
            }
        except Exception as e:
            raise Exception(f"Error loading email {file_path}: {e}")
    
    def batch_load(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of loaded documents
        """
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in [f'.{fmt}' for fmt in self.supported_formats]:
                try:
                    doc = self.load_document(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
                    continue
        
        return documents 